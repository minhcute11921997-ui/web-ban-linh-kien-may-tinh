from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT = r"D:\word\Bao_cao_du_an_TMĐT.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def add_table(doc, headers, rows, widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="0B2545")
        set_cell_shading(hdr[i], "E8EEF5")
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], str(value))
    if widths_cm:
        set_table_widths(table, widths_cm)
    for row in table.rows:
        for cell in row.cells:
            cell.margin_top = Cm(0.08)
            cell.margin_bottom = Cm(0.08)
            cell.margin_left = Cm(0.12)
            cell.margin_right = Cm(0.12)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181) if level <= 2 else RGBColor(31, 77, 120)
    return p


def add_para(doc, text="", style=None, bold_label=None):
    p = doc.add_paragraph(style=style)
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
    if text:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    return p


def add_number(doc, text):
    add_number.counter += 1
    p = doc.add_paragraph()
    r = p.add_run(f"{add_number.counter}. ")
    r.bold = False
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p


add_number.counter = 0


def reset_numbers():
    add_number.counter = 0


def add_page_break(doc):
    doc.add_page_break()


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(12 if name != "Heading 3" else 8)
        st.paragraph_format.space_after = Pt(6 if name != "Heading 3" else 4)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Báo cáo bài tập lớn - Phát triển hệ thống TMĐT").font.size = Pt(9)


def add_cover(doc):
    for text, size, bold in [
        ("HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG", 14, True),
        ("KHOA CÔNG NGHỆ THÔNG TIN 1", 13, True),
        ("oOo", 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÀI TẬP LỚN")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MÔN PHÁT TRIỂN HỆ THỐNG TMĐT")
    r.bold = True
    r.font.size = Pt(15)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Đề tài: Website thương mại điện tử bán linh kiện máy tính")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(46, 116, 181)

    doc.add_paragraph()
    doc.add_paragraph()
    for label in [
        "LỚP: ........................................................",
        "Số thứ tự nhóm: ...............................................",
        "GVHD: TS. Lê Văn Vịnh",
        "SVTH: .........................................................",
        "      .........................................................",
        "      .........................................................",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(3.0)
        r = p.add_run(label)
        r.font.size = Pt(12)

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HÀ NỘI, Tháng 05, Năm 2026")
    r.bold = True
    r.font.size = Pt(12)
    add_page_break(doc)


def add_toc(doc):
    add_heading(doc, "MỤC LỤC", 1)
    entries = [
        "Chương I. Mô tả, khảo sát và xác định yêu cầu bài toán",
        "Chương II. Kiến thức áp dụng",
        "Chương III. Phân tích thiết kế hệ thống",
        "Chương IV. Cài đặt và hướng dẫn sử dụng",
        "Kết luận",
        "Hướng phát triển",
        "Tài liệu tham khảo",
        "Phân công công việc",
    ]
    for e in entries:
        add_para(doc, e)
    add_page_break(doc)


def add_chapter_1(doc):
    add_heading(doc, "Chương I. MÔ TẢ, KHẢO SÁT VÀ XÁC ĐỊNH YÊU CẦU BÀI TOÁN", 1)
    add_heading(doc, "1.1 Mô tả bài toán", 2)
    add_para(
        doc,
        "Dự án xây dựng một website thương mại điện tử phục vụ bán linh kiện máy tính và sản phẩm công nghệ. "
        "Hệ thống cho phép khách hàng xem danh mục, tìm kiếm, lọc sản phẩm theo hãng/thông số/khoảng giá, xem chi tiết, "
        "thêm vào giỏ hàng, đặt hàng, thanh toán COD hoặc VNPay, quản lý địa chỉ giao hàng và theo dõi đơn hàng. "
        "Đối với quản trị viên, hệ thống cung cấp khu vực admin để quản lý sản phẩm, danh mục, đơn hàng, người dùng, mã giảm giá, doanh thu, banner và chương trình flash sale."
    )
    add_para(
        doc,
        "Mục tiêu của bài toán là tạo một hệ thống có đầy đủ luồng mua hàng cơ bản, có phân quyền người dùng, có khả năng mở rộng và dễ vận hành trong môi trường web hiện đại."
    )

    add_heading(doc, "1.2 Khảo sát, xác định yêu cầu bài toán", 2)
    add_heading(doc, "1.2.1 Đối tượng sử dụng", 3)
    add_bullet(doc, "Khách vãng lai: xem danh sách sản phẩm, lọc/tìm kiếm, xem chi tiết và đánh giá sản phẩm.")
    add_bullet(doc, "Khách hàng đã đăng nhập: quản lý hồ sơ, địa chỉ, giỏ hàng, đặt hàng, thanh toán, theo dõi và hủy đơn khi phù hợp.")
    add_bullet(doc, "Quản trị viên: quản lý dữ liệu sản phẩm, danh mục, đơn hàng, người dùng, mã giảm giá, banner, doanh thu và flash sale.")

    add_heading(doc, "1.2.2 Yêu cầu chức năng", 3)
    rows = [
        ("Đăng ký/đăng nhập", "Tạo tài khoản, đăng nhập bằng tài khoản thường hoặc Google, cấp access token và refresh token."),
        ("Trang chủ", "Hiển thị banner động, sản phẩm nổi bật, flash sale, danh mục, bộ lọc và danh sách sản phẩm."),
        ("Sản phẩm", "Xem danh sách, tìm kiếm, lọc theo danh mục/hãng/thông số/giá, xem chi tiết và đánh giá."),
        ("Giỏ hàng", "Thêm, sửa số lượng, xóa từng sản phẩm hoặc xóa toàn bộ giỏ hàng; kiểm tra tồn kho."),
        ("Thanh toán", "Tính tổng tiền, phí vận chuyển, áp dụng mã giảm giá, tạo đơn hàng và hỗ trợ COD/VNPay."),
        ("Đơn hàng", "Theo dõi trạng thái đơn, xem chi tiết, hủy đơn trong điều kiện cho phép."),
        ("Admin", "CRUD sản phẩm, danh mục, người dùng, đơn hàng, mã giảm giá; thống kê dashboard/doanh thu; quản lý banner."),
    ]
    add_table(doc, ["Nhóm chức năng", "Mô tả"], rows, [4.0, 12.0])

    add_heading(doc, "1.2.3 Yêu cầu phi chức năng", 3)
    add_bullet(doc, "Bảo mật: xác thực JWT, refresh token lưu trong CSDL, phân quyền admin, kiểm tra dữ liệu đầu vào bằng Joi.")
    add_bullet(doc, "Tính đúng đắn: kiểm tra tồn kho trước khi thêm giỏ hàng/đặt hàng, kiểm tra mã giảm giá và trạng thái thanh toán.")
    add_bullet(doc, "Khả dụng: giao diện React có phân trang, lọc, thông báo toast, loading state và route bảo vệ.")
    add_bullet(doc, "Dễ bảo trì: backend tách controller, route, middleware; frontend tách page, component, store và API client.")
    add_bullet(doc, "Mở rộng: có thể bổ sung chatbot, liên hệ, thông báo, báo cáo nâng cao và tích hợp vận chuyển thực tế.")


def add_chapter_2(doc):
    add_heading(doc, "Chương II. KIẾN THỨC ÁP DỤNG", 1)
    add_heading(doc, "2.1 Phân tích và thiết kế hệ thống", 2)
    add_para(doc, "Nhóm sử dụng cách tiếp cận phân tích chức năng theo vai trò người dùng, từ đó xác định các module chính và luồng dữ liệu giữa giao diện, API và cơ sở dữ liệu.")
    add_heading(doc, "2.1.1 Biểu đồ phân cấp chức năng BFD", 3)
    add_bullet(doc, "Hệ thống khách hàng: xác thực, xem sản phẩm, giỏ hàng, checkout, thanh toán, đơn hàng, hồ sơ, đánh giá.")
    add_bullet(doc, "Hệ thống quản trị: dashboard, sản phẩm, danh mục, đơn hàng, doanh thu, người dùng, mã giảm giá, banner.")
    add_bullet(doc, "Hạ tầng dùng chung: xác thực JWT, upload ảnh, static file, middleware phân quyền, kết nối MySQL.")
    add_heading(doc, "2.1.2 Biểu đồ luồng dữ liệu DFD", 3)
    reset_numbers()
    add_number(doc, "Người dùng gửi thao tác từ React UI thông qua Axios/fetch.")
    add_number(doc, "Express route nhận request, middleware xác thực token nếu route yêu cầu đăng nhập.")
    add_number(doc, "Controller xử lý nghiệp vụ, truy vấn MySQL hoặc thao tác file upload.")
    add_number(doc, "Backend trả JSON về frontend, frontend cập nhật state và hiển thị kết quả.")

    add_heading(doc, "2.2 Quản trị hệ thống", 2)
    add_para(doc, "Khu vực quản trị được bảo vệ bằng PrivateRoute với điều kiện tài khoản có role admin. Menu admin gồm Dashboard, Sản phẩm, Đơn hàng, Doanh thu, Banner, Danh mục, Người dùng và Mã giảm giá.")
    add_para(doc, "Các API quản trị dùng middleware verifyToken và verifyAdmin để hạn chế truy cập. Ảnh sản phẩm và banner được upload qua multer, lưu trong thư mục uploads và phục vụ qua static route /uploads.")

    add_heading(doc, "2.3 Cơ sở dữ liệu", 2)
    add_para(doc, "Hệ thống sử dụng MySQL, truy cập qua mysql2/promise. Dữ liệu chính được tổ chức theo các nhóm bảng: người dùng, sản phẩm, giỏ hàng, đơn hàng, thanh toán, mã giảm giá, đánh giá và địa chỉ giao hàng.")
    rows = [
        ("users", "Thông tin tài khoản, hồ sơ và vai trò customer/admin."),
        ("refresh_tokens", "Refresh token còn hiệu lực để chống replay và hỗ trợ đăng xuất."),
        ("products", "Thông tin sản phẩm, giá, tồn kho, ảnh, danh mục, flash sale, trạng thái hiển thị."),
        ("categories", "Danh mục sản phẩm."),
        ("product_specifications", "Thông số kỹ thuật theo từng sản phẩm để lọc nâng cao."),
        ("cart, cart_items", "Giỏ hàng và từng sản phẩm trong giỏ."),
        ("orders, order_items", "Đơn hàng, chi tiết đơn, trạng thái xử lý và tổng tiền."),
        ("payments", "Thông tin thanh toán COD/VNPay và trạng thái thanh toán."),
        ("discounts", "Mã giảm giá, điều kiện áp dụng, số lượt dùng, thời hạn."),
        ("reviews", "Đánh giá sản phẩm của khách hàng."),
        ("addresses", "Địa chỉ giao hàng của người dùng."),
    ]
    add_table(doc, ["Bảng/nhóm bảng", "Vai trò"], rows, [4.5, 11.5])

    add_heading(doc, "2.4 Ngôn ngữ lập trình và công nghệ", 2)
    rows = [
        ("Frontend", "React 19, Vite, Tailwind CSS, React Router, Zustand, Axios, Lucide React, React Toastify."),
        ("Backend", "Node.js, Express 5, CommonJS, MySQL2, JSON Web Token, Joi, Multer, Google Auth Library, VNPay SDK."),
        ("CSDL", "MySQL."),
        ("Thanh toán", "COD và VNPay sandbox/callback/IPN."),
        ("Xác thực", "JWT access token, refresh token, Google OAuth Client ID."),
        ("Upload", "Ảnh sản phẩm và banner lưu local trong thư mục uploads."),
    ]
    add_table(doc, ["Thành phần", "Công nghệ sử dụng"], rows, [4.0, 12.0])


def add_chapter_3(doc):
    add_heading(doc, "Chương III. PHÂN TÍCH THIẾT KẾ HỆ THỐNG", 1)
    add_heading(doc, "3.1 Phân tích thiết kế CSDL", 2)
    add_para(doc, "CSDL được thiết kế theo hướng tách các thực thể chính, giảm lặp dữ liệu và hỗ trợ mở rộng nghiệp vụ. Quan hệ tiêu biểu gồm: một người dùng có nhiều địa chỉ, một người dùng có một giỏ hàng, một giỏ hàng có nhiều cart_items, một đơn hàng có nhiều order_items, một sản phẩm thuộc một danh mục và có nhiều thông số kỹ thuật.")
    rows = [
        ("users - orders", "1 - N", "Một người dùng có thể tạo nhiều đơn hàng."),
        ("users - addresses", "1 - N", "Một người dùng lưu nhiều địa chỉ giao hàng."),
        ("users - cart", "1 - 1", "Mỗi người dùng có một giỏ hàng đang hoạt động."),
        ("cart - cart_items", "1 - N", "Một giỏ hàng gồm nhiều dòng sản phẩm."),
        ("products - cart_items", "1 - N", "Một sản phẩm có thể xuất hiện trong nhiều giỏ hàng."),
        ("orders - order_items", "1 - N", "Một đơn hàng gồm nhiều sản phẩm."),
        ("products - reviews", "1 - N", "Một sản phẩm có nhiều đánh giá."),
        ("categories - products", "1 - N", "Một danh mục gồm nhiều sản phẩm."),
    ]
    add_table(doc, ["Quan hệ", "Kiểu", "Ý nghĩa"], rows, [4.0, 2.2, 9.8])

    add_heading(doc, "3.2 Phân tích thiết kế chức năng", 2)
    add_heading(doc, "3.2.1 Module khách hàng", 3)
    add_bullet(doc, "Trang chủ: banner động, sản phẩm nổi bật, flash sale, danh mục, tìm kiếm và bộ lọc.")
    add_bullet(doc, "Sản phẩm: xem danh sách, chi tiết, thông số, ảnh, tồn kho, giá khuyến mãi và đánh giá.")
    add_bullet(doc, "Tài khoản: đăng ký, đăng nhập thường, đăng nhập Google, refresh token, đăng xuất, cập nhật hồ sơ.")
    add_bullet(doc, "Giỏ hàng/checkout: kiểm tra tồn kho, áp mã giảm giá, tính phí vận chuyển, tạo đơn và thanh toán.")
    add_bullet(doc, "Đơn hàng: xem danh sách, chi tiết, trạng thái thanh toán và hủy đơn.")

    add_heading(doc, "3.2.2 Module quản trị", 3)
    add_bullet(doc, "Dashboard: tổng quan đơn hàng, doanh thu, người dùng và sản phẩm.")
    add_bullet(doc, "Quản lý sản phẩm: thêm/sửa/xóa, upload ảnh, bật/tắt hiển thị, thiết lập flash sale.")
    add_bullet(doc, "Quản lý đơn hàng: lọc, xem chi tiết, cập nhật trạng thái, xóa/hủy theo nghiệp vụ.")
    add_bullet(doc, "Quản lý doanh thu: báo cáo theo ngày/tháng.")
    add_bullet(doc, "Quản lý banner: upload ảnh từ máy, chọn vị trí banner 1/2/3 trên trang chủ và cập nhật link.")
    add_bullet(doc, "Quản lý danh mục, người dùng và mã giảm giá.")

    add_heading(doc, "3.2.3 Thiết kế API chính", 3)
    rows = [
        ("Auth", "/api/auth/register, /login, /google, /refresh, /logout, /profile", "Tài khoản và xác thực."),
        ("Products", "/api/products, /featured, /on-sale, /filters/:categoryId, /:id/specs", "Sản phẩm, lọc, sale, thông số."),
        ("Cart", "/api/cart, /add, /item/:id, /clear", "Quản lý giỏ hàng."),
        ("Checkout/Order", "/api/checkout/calculate, /api/orders", "Tính tiền và tạo/quản lý đơn hàng."),
        ("Payments", "/api/payments/create-order, /vnpay-callback, /vnpay-ipn", "VNPay và trạng thái thanh toán."),
        ("Discounts", "/api/discounts/validate, /available, /admin/*", "Mã giảm giá khách hàng và admin."),
        ("Admin", "/api/dashboard, /api/users, /api/categories", "Thống kê, người dùng, danh mục."),
        ("Banners", "/api/banners, /api/banners/:position", "Lấy và cập nhật banner trang chủ."),
    ]
    add_table(doc, ["Module", "Endpoint tiêu biểu", "Mục đích"], rows, [3.0, 7.0, 6.0])

    add_heading(doc, "3.3 Các chức năng chưa làm được", 2)
    add_bullet(doc, "Chưa có chatbot tư vấn tự động và trang liên hệ/giải đáp thắc mắc đầy đủ như mẫu gợi ý.")
    add_bullet(doc, "Chưa tích hợp đơn vị vận chuyển thật; phí vận chuyển mới ở mức tính toán nội bộ.")
    add_bullet(doc, "Chưa có hệ thống thông báo realtime cho đơn hàng và tồn kho.")
    add_bullet(doc, "Chưa có phân quyền chi tiết theo từng nhóm admin; hiện dùng phân quyền customer/admin.")
    add_bullet(doc, "Chưa có báo cáo tài chính nâng cao và xuất file Excel/PDF cho admin.")


def add_chapter_4(doc):
    add_heading(doc, "Chương IV. CÀI ĐẶT VÀ HƯỚNG DẪN SỬ DỤNG", 1)
    add_heading(doc, "4.1 Cài đặt CSDL", 2)
    reset_numbers()
    add_number(doc, "Cài đặt MySQL và tạo database cho dự án.")
    add_number(doc, "Cấu hình thông tin DB_HOST, DB_PORT, DB_USER, DB_PASSWORD, DB_NAME trong be/.env.")
    add_number(doc, "Đảm bảo các bảng nghiệp vụ như users, products, categories, cart, orders, payments, discounts, reviews, addresses và refresh_tokens đã được tạo đúng cấu trúc.")
    add_number(doc, "Chạy backend để kiểm tra kết nối MySQL và các API.")

    add_heading(doc, "4.2 Cài đặt giả lập môi trường server hosting", 2)
    add_table(
        doc,
        ["Bước", "Lệnh/cấu hình"],
        [
            ("1", "cd be && npm install"),
            ("2", "Tạo be/.env gồm DB, JWT_SECRET, JWT_REFRESH_SECRET, ALLOWED_ORIGINS, FRONTEND_URL, GOOGLE_CLIENT_ID, VNPAY_* và PORT=3000."),
            ("3", "npm run start hoặc npm run dev để chạy backend tại http://localhost:3000."),
            ("4", "cd fe && npm install"),
            ("5", "Tạo fe/.env gồm VITE_API_URL=http://localhost:3000 và VITE_GOOGLE_CLIENT_ID."),
            ("6", "npm run dev để chạy frontend tại http://localhost:5173."),
        ],
        [2.0, 14.0],
    )

    add_heading(doc, "4.3 Giao diện User", 2)
    user_pages = [
        ("4.3.1 Giao diện trang chủ", "Hiển thị banner, sản phẩm nổi bật, flash sale, danh mục, tìm kiếm, lọc giá và sắp xếp."),
        ("4.3.2 Giao diện trang sản phẩm", "Danh sách sản phẩm dạng lưới, hỗ trợ lọc theo danh mục, hãng, thông số kỹ thuật và khoảng giá."),
        ("4.3.3 Giao diện trang chi tiết sản phẩm", "Ảnh, tên, hãng, giá, tồn kho, số lượng, thêm giỏ, mua ngay, thông số và đánh giá."),
        ("4.3.4 Giao diện trang tìm kiếm", "Tìm kiếm theo từ khóa, kết hợp với bộ lọc và sắp xếp."),
        ("4.3.5 Giao diện trang đăng ký tài khoản", "Đăng ký bằng form thông tin cá nhân hoặc đăng ký/đăng nhập bằng Google."),
        ("4.3.6 Giao diện trang đăng nhập", "Đăng nhập bằng username/email và mật khẩu hoặc Google OAuth."),
        ("4.3.7 Giao diện trang giỏ hàng", "Xem sản phẩm trong giỏ, cập nhật số lượng, xóa sản phẩm và chuyển sang checkout."),
        ("4.3.8 Giao diện trang xác nhận đặt hàng", "Chọn địa chỉ, mã giảm giá, phương thức thanh toán và xác nhận đơn."),
        ("4.3.9 Giao diện trang thanh toán", "Hỗ trợ COD và VNPay; VNPay có callback và trang kết quả thanh toán."),
        ("4.3.10 Giao diện trang đơn hàng", "Danh sách đơn theo trạng thái, xem chi tiết và hủy đơn khi được phép."),
        ("4.3.11 Giao diện trang hồ sơ", "Cập nhật thông tin cá nhân và quản lý địa chỉ giao hàng."),
    ]
    for title, desc in user_pages:
        add_heading(doc, title, 3)
        add_para(doc, desc)

    add_heading(doc, "4.4 Giao diện Admin", 2)
    admin_pages = [
        ("4.4.1 Giao diện trang đăng nhập phần Admin", "Admin đăng nhập bằng tài khoản có role admin, sau đó được PrivateRoute cho phép vào /admin."),
        ("4.4.2 Giao diện trang chủ Admin", "Dashboard tổng quan các chỉ số vận hành."),
        ("4.4.3 Giao diện trang thêm mới sản phẩm", "Form thêm sản phẩm, chọn danh mục, giá, tồn kho, ảnh URL hoặc upload từ máy."),
        ("4.4.4 Giao diện trang sửa sản phẩm", "Cập nhật thông tin sản phẩm, ảnh, trạng thái hiển thị và flash sale."),
        ("4.4.5 Giao diện trang xóa sản phẩm", "Xóa sản phẩm khỏi hệ thống theo quyền admin."),
        ("4.4.6 Giao diện trang quản lý tài khoản người dùng", "Xem, sửa vai trò/thông tin và xóa người dùng."),
        ("4.4.7 Giao diện trang thống kê", "Theo dõi doanh thu theo ngày/tháng và số liệu đơn hàng."),
        ("4.4.8 Giao diện trang quản lý banner", "Chọn tệp ảnh từ máy để thay Banner 1, Banner 2 hoặc Banner 3 trên trang chủ, kèm link điều hướng."),
        ("4.4.9 Giao diện trang quản lý danh mục", "Thêm, sửa, xóa danh mục sản phẩm."),
        ("4.4.10 Giao diện trang quản lý Sales", "Thiết lập flash sale, số lượng sale, phần trăm giảm và thời gian hiệu lực."),
        ("4.4.11 Giao diện trang quản lý mã giảm giá", "Tạo, sửa, bật/tắt và xóa mã giảm giá."),
    ]
    for title, desc in admin_pages:
        add_heading(doc, title, 3)
        add_para(doc, desc)


def add_conclusion(doc):
    add_heading(doc, "KẾT LUẬN", 1)
    add_para(doc, "Dự án đã xây dựng được một website thương mại điện tử tương đối đầy đủ cho quy trình bán hàng trực tuyến: từ xem sản phẩm, tìm kiếm, giỏ hàng, đặt hàng, thanh toán đến quản trị sản phẩm và đơn hàng. Hệ thống đã có phân quyền người dùng, đăng nhập Google, thanh toán VNPay, quản lý banner và các chức năng quản trị cần thiết.")
    add_para(doc, "Trong quá trình thực hiện, nhóm đã vận dụng kiến thức phân tích thiết kế hệ thống, thiết kế CSDL quan hệ, xây dựng REST API, xử lý xác thực bằng JWT, quản lý state phía frontend và tổ chức giao diện React theo component/page.")

    add_heading(doc, "HƯỚNG PHÁT TRIỂN", 1)
    add_bullet(doc, "Bổ sung chatbot tư vấn sản phẩm và trang giải đáp thắc mắc/liên hệ.")
    add_bullet(doc, "Tích hợp dịch vụ vận chuyển thật và theo dõi trạng thái giao hàng.")
    add_bullet(doc, "Bổ sung thông báo realtime cho khách hàng và admin.")
    add_bullet(doc, "Hoàn thiện phân quyền chi tiết cho nhiều nhóm quản trị.")
    add_bullet(doc, "Tối ưu hiệu năng frontend bằng code splitting và tối ưu kích thước ảnh/banner.")
    add_bullet(doc, "Bổ sung test tự động cho API, checkout, thanh toán và phân quyền.")

    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "React Documentation: https://react.dev/",
        "Vite Documentation: https://vite.dev/",
        "Express Documentation: https://expressjs.com/",
        "MySQL Documentation: https://dev.mysql.com/doc/",
        "JSON Web Token: https://jwt.io/",
        "Google Identity Services: https://developers.google.com/identity",
        "VNPay Sandbox/Developer Documentation.",
    ]
    reset_numbers()
    for ref in refs:
        add_number(doc, ref)

    add_heading(doc, "PHÂN CÔNG CÔNG VIỆC", 1)
    add_table(
        doc,
        ["STT", "Mã SV", "Họ và tên SV", "Công việc", "Xác nhận"],
        [
            ("1", "", "Nhóm trưởng", "Phân tích yêu cầu, thiết kế chức năng, tổng hợp báo cáo.", ""),
            ("2", "", "", "Xây dựng backend: auth, sản phẩm, giỏ hàng, đơn hàng, thanh toán, admin API.", ""),
            ("3", "", "", "Xây dựng frontend: giao diện user, admin, checkout, Google Auth, banner.", ""),
            ("4", "", "", "Kiểm thử, sửa lỗi, hoàn thiện tài liệu và hướng dẫn cài đặt.", ""),
        ],
        [1.2, 2.5, 4.0, 6.8, 1.5],
    )


def main():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc)
    add_chapter_4(doc)
    add_conclusion(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
