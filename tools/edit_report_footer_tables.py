from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def clear_cell_shading(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is not None:
        tc_pr.remove(shd)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def main():
    matches = list(Path("D:/word").glob("Bao_cao_du_an*.docx"))
    if not matches:
        raise FileNotFoundError("Khong tim thay file Bao_cao_du_an*.docx trong D:/word")

    path = matches[0]
    doc = Document(str(path))

    for section in doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            paragraph.clear()
        first = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        first.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_field(first)

    for table in doc.tables:
        if not table.rows:
            continue
        for cell in table.rows[0].cells:
            clear_cell_shading(cell)

    doc.save(str(path))
    print(path)


if __name__ == "__main__":
    main()
